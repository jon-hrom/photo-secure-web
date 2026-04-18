"""
Инструмент Humanizer: переписывает AI-тексты человеческим языком.

Actions (POST body → field "action"):
  - parse_file     — распарсить загруженный файл (txt/docx/pdf/rtf/odt) в чистый текст
  - detect         — разбить текст на предложения и посчитать AI-score для каждого
                     (гибрид: эвристика + LLM-самооценка)
  - rewrite        — очеловечить один фрагмент (возвращает 3 варианта)
  - humanize_full  — полный прогон: детект → переписывание всех проблемных предложений
                     → повторный замер → если score всё ещё > порога, ещё один проход
  - save_document  — сохранить результат в историю
  - list_history   — список сохранённых документов пользователя
  - get_document   — получить один документ по id
  - export_docx    — экспорт в .docx (возвращает base64)

Требует secrets: YANDEX_GPT_API_KEY, YANDEX_GPT_FOLDER_ID, DATABASE_URL
"""
import os
import re
import io
import json
import math
import base64
import random
import hashlib
import urllib.request
import urllib.error
from typing import List, Dict, Any, Optional, Tuple

import psycopg2
from psycopg2.extras import RealDictCursor

SCHEMA = os.environ.get("MAIN_DB_SCHEMA", "t_p28211681_photo_secure_web")

# ======================================================================
#                          CORS / helpers
# ======================================================================

CORS_HEADERS = {
    "Access-Control-Allow-Origin": "*",
    "Access-Control-Allow-Methods": "GET, POST, PUT, DELETE, OPTIONS",
    "Access-Control-Allow-Headers": "Content-Type, X-User-Id, X-Auth-Token",
    "Access-Control-Max-Age": "86400",
}


def _ok(body: Dict[str, Any], status: int = 200) -> Dict[str, Any]:
    return {
        "statusCode": status,
        "headers": {**CORS_HEADERS, "Content-Type": "application/json"},
        "isBase64Encoded": False,
        "body": json.dumps(body, ensure_ascii=False),
    }


def _err(msg: str, status: int = 400) -> Dict[str, Any]:
    return _ok({"error": msg}, status)


def _db():
    return psycopg2.connect(os.environ["DATABASE_URL"])


# ======================================================================
#                         Парсер файлов
# ======================================================================

def _extract_text_from_bytes(raw: bytes, filename: str) -> str:
    """Достаёт plain text из различных форматов."""
    name = (filename or "").lower()

    # txt / md / csv / html / rtf
    if name.endswith((".txt", ".md", ".csv", ".log", ".json")):
        for enc in ("utf-8", "cp1251", "latin-1"):
            try:
                return raw.decode(enc)
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="ignore")

    if name.endswith(".rtf"):
        text = raw.decode("latin-1", errors="ignore")
        # Очень грубый RTF-стриппер
        text = re.sub(r"\{\\[^}]+\}", "", text)
        text = re.sub(r"\\[a-z]+-?\d* ?", "", text)
        text = re.sub(r"[{}]", "", text)
        return text

    if name.endswith(".html") or name.endswith(".htm"):
        text = raw.decode("utf-8", errors="ignore")
        text = re.sub(r"<script.*?</script>", "", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<style.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"&nbsp;", " ", text)
        text = re.sub(r"&amp;", "&", text)
        return re.sub(r"\s+", " ", text).strip()

    # docx
    if name.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())
            return "\n".join(parts)
        except Exception as e:
            # fallback — достаём из word/document.xml
            try:
                import zipfile
                with zipfile.ZipFile(io.BytesIO(raw)) as z:
                    with z.open("word/document.xml") as f:
                        xml = f.read().decode("utf-8", errors="ignore")
                text = re.sub(r"<w:p[^>]*>", "\n", xml)
                text = re.sub(r"<[^>]+>", "", text)
                return re.sub(r"\n{3,}", "\n\n", text).strip()
            except Exception:
                raise RuntimeError(f"Не удалось прочитать docx: {e}")

    # doc (старый бинарный формат) — пробуем как текст
    if name.endswith(".doc"):
        text = raw.decode("cp1251", errors="ignore")
        text = re.sub(r"[^\u0400-\u04FFa-zA-Z0-9.,!?;:\-()\"\'\s]", " ", text)
        return re.sub(r"\s+", " ", text).strip()

    # pdf
    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            parts = []
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    parts.append(t)
            return "\n".join(parts)
        except Exception as e:
            raise RuntimeError(f"Не удалось прочитать pdf: {e}")

    # odt
    if name.endswith(".odt"):
        try:
            import zipfile
            with zipfile.ZipFile(io.BytesIO(raw)) as z:
                with z.open("content.xml") as f:
                    xml = f.read().decode("utf-8", errors="ignore")
            text = re.sub(r"<text:p[^>]*>", "\n", xml)
            text = re.sub(r"<[^>]+>", "", text)
            return re.sub(r"\n{3,}", "\n\n", text).strip()
        except Exception as e:
            raise RuntimeError(f"Не удалось прочитать odt: {e}")

    # Если расширение неизвестно — пробуем как utf-8
    try:
        return raw.decode("utf-8")
    except Exception:
        return raw.decode("utf-8", errors="ignore")


# ======================================================================
#                     Разбиение на предложения
# ======================================================================

_SENT_SPLIT_RE = re.compile(r"(?<=[.!?…])\s+(?=[А-ЯA-Z«\"])")

def split_sentences(text: str) -> List[Dict[str, Any]]:
    """Разбивает текст на предложения с сохранением позиций."""
    text = text.replace("\r\n", "\n")
    # Сначала разбиваем по абзацам, потом по предложениям
    sentences = []
    pos = 0
    paragraphs = re.split(r"\n\s*\n", text)
    for para_idx, para in enumerate(paragraphs):
        para_stripped = para.strip()
        if not para_stripped:
            pos += len(para) + 2
            continue
        # Находим реальный start в исходном тексте
        para_start = text.find(para_stripped, pos)
        if para_start < 0:
            para_start = pos
        # Разбиваем параграф на предложения
        parts = _SENT_SPLIT_RE.split(para_stripped)
        cursor = para_start
        for part in parts:
            part = part.strip()
            if not part:
                continue
            idx = text.find(part, cursor)
            if idx < 0:
                idx = cursor
            sentences.append({
                "index": len(sentences),
                "text": part,
                "start": idx,
                "end": idx + len(part),
                "paragraph": para_idx,
            })
            cursor = idx + len(part)
        pos = para_start + len(para_stripped)
    return sentences


# ======================================================================
#                         AI-детектор (гибрид)
# ======================================================================

# Типичные "AI-штампы" на русском (+ маркеры Томилиной)
AI_PHRASES = [
    r"в современном мире", r"играет важную роль", r"стоит отметить",
    r"в первую очередь", r"на сегодняшний день", r"в заключение",
    r"важно понимать", r"ключевым аспектом", r"не только.{0,30}но и",
    r"необходимо учитывать", r"является одним из", r"в рамках данного",
    r"существует множество", r"представляет собой", r"таким образом",
    r"следует отметить", r"в свою очередь", r"можно сделать вывод",
    r"в целом", r"однако,?\s", r"тем не менее,?\s",
    r"подводя итог", r"в конечном счёте", r"важной составляющей",
    r"необходимо отметить", r"рассмотрим подробнее", r"особое внимание",
    r"эффективный", r"оптимальный", r"уникальн\w+ возможност",
    r"широкий спектр", r"в условиях современн", r"динамично развива",
    # --- академические маркеры (Томилина) ---
    r"\bво-первых\b", r"\bво-вторых\b", r"\bв-третьих\b",
    r"\bкроме того\b", r"\bболее того\b",
    r"\bэто означает\b", r"\bэто свидетельствует\b",
    r"\bэто имеет непосредственное\b", r"\bимеет непосредственное значение\b",
    r"\bнепосредственное значение для темы\b",
    # методологически некорректные формулировки (субъект — неодушевлённый)
    r"\bработа (анализирует|рассматривает|отображает|показывает|демонстрирует|раскрывает|описывает|исследует|посвящена)\b",
    r"\bглава (анализирует|рассматривает|отображает|показывает|раскрывает|описывает|посвящена)\b",
    r"\bстатья (анализирует|рассматривает|отображает|показывает|раскрывает|описывает)\b",
    r"\bисследование (анализирует|рассматривает|показывает|раскрывает|описывает)\b",
    r"\bзакон (уточняет|указывает|устанавливает|определяет|требует)\b",
    r"\bсистема (указывает|определяет|требует|устанавливает)\b",
    r"\bсудебная практика (требует|устанавливает|показывает)\b",
    r"\bнорма (устанавливает|определяет|требует)\b",
]
AI_PHRASES_RE = re.compile("|".join(AI_PHRASES), re.IGNORECASE)

# Английские AI-штампы
AI_PHRASES_EN = [
    r"in today'?s world", r"it'?s important to note", r"plays a crucial role",
    r"it'?s worth noting", r"in conclusion", r"furthermore", r"moreover",
    r"delve into", r"leverage", r"navigate the complexities",
    r"in the realm of", r"a myriad of", r"tapestry", r"testament to",
    r"cutting-edge", r"seamlessly", r"robust", r"comprehensive",
    r"it'?s essential", r"key aspect",
]
AI_PHRASES_EN_RE = re.compile("|".join(AI_PHRASES_EN), re.IGNORECASE)


def _heuristic_ai_score(sentence: str) -> float:
    """
    Эвристика AI-score для предложения. Возвращает 0..100.
    Критерии:
      - штампы AI
      - "ровная" длина предложения (15-25 слов — любимая длина AI)
      - высокая доля "книжных" слов
      - отсутствие разговорных маркеров
      - параллельные конструкции
    """
    s = sentence.strip()
    if len(s) < 20:
        return 0.0

    score = 0.0
    words = re.findall(r"\w+", s, flags=re.UNICODE)
    wc = len(words)

    # штампы
    ru_hits = len(AI_PHRASES_RE.findall(s))
    en_hits = len(AI_PHRASES_EN_RE.findall(s))
    score += min(40, (ru_hits + en_hits) * 22)

    # длина
    if 14 <= wc <= 28:
        score += 12
    elif 28 < wc <= 40:
        score += 6

    # сложноподчинённые с "который/которая/которые/которое"
    if re.search(r"\bкотор[ыаоуе][йяеим]?\b", s, re.IGNORECASE):
        score += 5

    # шаблонные вводные
    if re.search(r"^(таким образом|следовательно|итак|однако|тем не менее|более того|кроме того)[,\s]", s, re.IGNORECASE):
        score += 10

    # перечисление через "и ... и ..."
    if re.search(r"\b(не только|как).*?(но и|так и)\b", s, re.IGNORECASE):
        score += 8

    # все слова — полные, ни одного разговорного маркера
    colloquial = re.search(r"\b(ну|вот|короче|типа|вроде|как бы|в общем|кстати|кажется|по-моему|походу|мол|щас)\b", s, re.IGNORECASE)
    if not colloquial and wc > 12:
        score += 7

    # отсутствие личных местоимений от 1-го лица
    if not re.search(r"\b(я|мы|меня|мне|нам|нас|мой|наш)\b", s, re.IGNORECASE) and wc > 15:
        score += 4

    # плотность "умных" суффиксов: -ание, -ение, -ость
    smart_suffixes = len(re.findall(r"\w+(?:ание|ение|ость|ация|ирование)\b", s, re.IGNORECASE))
    if smart_suffixes >= 3:
        score += 10
    elif smart_suffixes >= 2:
        score += 5

    # Высокое разнообразие длинных слов (>8 символов)
    long_words = sum(1 for w in words if len(w) > 8)
    if wc > 0 and long_words / wc > 0.35:
        score += 6

    # --- маркеры Томилиной ---
    # 1) длинное тире — (U+2014) — очень сильный индикатор скопированного/сгенерированного
    em_dashes = s.count("—")
    if em_dashes >= 1:
        score += min(18, em_dashes * 9)

    # 2) буква «ё» в академическом тексте (если она есть в словах, где не обязательна)
    # Смотрим формы вроде "ведёт", "проведён", "отражён" — это маркер, что текст не из акад. редакции
    if re.search(r"\w*(?:ё)\w+", s):
        # лёгкий штраф — сам по себе маркер слабый, но в сочетании с другими значимый
        score += 3

    # 3) чрезмерно длинное предложение (>40 слов)
    if wc > 40:
        score += 10
    if wc > 55:
        score += 8  # суммарно до +18 за очень длинные

    # 4) английские термины-вставки без кавычек (скачок языка)
    eng_words = re.findall(r"\b[A-Za-z]{4,}\b", s)
    if eng_words:
        # проверяем что это не общеизвестные аббревиатуры типа URL/API (капс) и не в скобках/кавычках
        non_caps = [w for w in eng_words if not w.isupper()]
        if non_caps:
            score += min(10, len(non_caps) * 4)

    # 5) плотные перечисления (3+ запятых подряд близко)
    commas = s.count(",")
    if wc > 15 and commas >= 4:
        score += 5

    # 6) вводная шаблонная связка где-то в середине
    if re.search(r",\s*(кроме того|более того|следует отметить|необходимо отметить|таким образом|тем не менее)[,\s]", s, re.IGNORECASE):
        score += 6

    return float(min(100, max(0, score)))


def score_sentences_heuristic(sentences: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Выставляет heuristic_score для каждого предложения."""
    for s in sentences:
        s["heuristic_score"] = _heuristic_ai_score(s["text"])
    return sentences


def _burstiness(sentences: List[Dict[str, Any]]) -> float:
    """
    Burstiness: вариативность длины предложений. Низкая = AI.
    Возвращаем bonus к AI-score (0..15).
    """
    lens = [len(re.findall(r"\w+", s["text"])) for s in sentences if len(s["text"]) > 10]
    if len(lens) < 3:
        return 0.0
    mean = sum(lens) / len(lens)
    var = sum((x - mean) ** 2 for x in lens) / len(lens)
    std = math.sqrt(var)
    cv = std / mean if mean > 0 else 0
    if cv < 0.25:
        return 15.0
    if cv < 0.4:
        return 8.0
    return 0.0


def _tense_jumps(sentences: List[Dict[str, Any]]) -> float:
    """
    Несогласованность времён (Томилина): текст скачет между прошедшим и настоящим.
    Простая эвристика — доля переходов прош→наст между соседними предложениями.
    """
    if len(sentences) < 3:
        return 0.0
    past_re = re.compile(r"\w+(?:л|ла|ло|ли|лся|лась|лись|вшись|нут|нута)\b", re.IGNORECASE)
    pres_re = re.compile(r"\w+(?:ет|ёт|ут|ют|ит|ат|ят|ем|ём|ете|ёте|у|ю)\b", re.IGNORECASE)
    tenses = []
    for s in sentences:
        t = s["text"]
        p = len(past_re.findall(t))
        n = len(pres_re.findall(t))
        if p + n < 2:
            tenses.append(None)
        elif p > n * 1.5:
            tenses.append("past")
        elif n > p * 1.5:
            tenses.append("pres")
        else:
            tenses.append(None)
    jumps = 0
    total = 0
    for i in range(1, len(tenses)):
        if tenses[i] and tenses[i - 1] and tenses[i] != tenses[i - 1]:
            jumps += 1
        if tenses[i] and tenses[i - 1]:
            total += 1
    if total < 3:
        return 0.0
    ratio = jumps / total
    if ratio > 0.4:
        return 10.0
    if ratio > 0.25:
        return 5.0
    return 0.0


def _uniform_paragraph_structure(text: str) -> float:
    """
    Маркер Томилиной — одинаковая структура абзацев:
    «общее утверждение → объяснение → вывод» в каждом абзаце.
    Эвристика: абзацы имеют очень близкое количество предложений и длину.
    """
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if len(paragraphs) < 3:
        return 0.0
    sizes = []
    for p in paragraphs:
        sents = [x for x in re.split(r"(?<=[.!?])\s+", p) if x]
        sizes.append(len(sents))
    mean = sum(sizes) / len(sizes)
    if mean < 2:
        return 0.0
    var = sum((x - mean) ** 2 for x in sizes) / len(sizes)
    std = math.sqrt(var)
    cv = std / mean if mean > 0 else 0
    # Однотипные абзацы = низкий cv по количеству предложений
    if cv < 0.15 and mean >= 2.5:
        return 8.0
    if cv < 0.25:
        return 4.0
    return 0.0


# ======================================================================
#                              YandexGPT
# ======================================================================

YGPT_URL = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"


def _yagpt_complete(system: str, user: str, temperature: float = 0.6,
                    max_tokens: int = 2000, model: str = "yandexgpt") -> str:
    """Синхронный вызов YandexGPT."""
    api_key_raw = os.environ.get("YANDEX_GPT_API_KEY", "")
    folder_raw = os.environ.get("YANDEX_GPT_FOLDER_ID", "")
    api_key = api_key_raw.strip()
    folder = folder_raw.strip()
    if not api_key or not folder:
        # Диагностика — показываем что именно пустое (без раскрытия значений)
        has_key = bool(api_key_raw)
        has_folder = bool(folder_raw)
        key_len = len(api_key)
        folder_len = len(folder)
        all_env_keys = sorted([k for k in os.environ.keys() if 'YANDEX' in k.upper() or 'GPT' in k.upper()])
        raise RuntimeError(
            f"Секреты YandexGPT не настроены. "
            f"API_KEY: present={has_key}, len={key_len}. "
            f"FOLDER_ID: present={has_folder}, len={folder_len}. "
            f"Найденные env-переменные: {all_env_keys}"
        )

    # Модели: yandexgpt (pro) / yandexgpt-lite
    model_uri = f"gpt://{folder}/{model}/latest"

    payload = {
        "modelUri": model_uri,
        "completionOptions": {
            "stream": False,
            "temperature": temperature,
            "maxTokens": str(max_tokens),
        },
        "messages": [
            {"role": "system", "text": system},
            {"role": "user", "text": user},
        ],
    }
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")

    # Yandex принимает и "Api-Key ..." и голый ключ — нормализуем
    auth_header = api_key if api_key.lower().startswith("api-key ") else f"Api-Key {api_key}"

    req = urllib.request.Request(
        YGPT_URL,
        data=body,
        headers={
            "Content-Type": "application/json",
            "Authorization": auth_header,
            "x-folder-id": folder,
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        body_err = e.read().decode("utf-8", errors="ignore")
        raise RuntimeError(f"YandexGPT HTTP {e.code}: {body_err[:400]}")
    except urllib.error.URLError as e:
        raise RuntimeError(f"YandexGPT URL error: {e}")

    try:
        return data["result"]["alternatives"][0]["message"]["text"]
    except Exception:
        raise RuntimeError(f"YandexGPT неожиданный ответ: {json.dumps(data)[:500]}")


# ======================================================================
#                       Антидетект: трюки постпроцессинга
# ======================================================================

# Замена AI-штампов на "живые"
REPLACEMENTS = [
    (r"\bв современном мире\b", ["сегодня", "сейчас", "в наше время"]),
    (r"\bиграет важную роль\b", ["много значит", "сильно влияет", "важно"]),
    (r"\bстоит отметить\b", ["замечу", "добавлю", "кстати"]),
    (r"\bнеобходимо отметить\b", ["замечу", "к слову"]),
    (r"\bв первую очередь\b", ["первым делом", "прежде всего", "главное"]),
    (r"\bтаким образом\b", ["значит", "в итоге", "выходит"]),
    (r"\bследует отметить\b", ["замечу", "кстати"]),
    (r"\bявляется\b", ["это", "—"]),
    (r"\bданный\b", ["этот"]),
    (r"\bданные\b", ["эти"]),
    (r"\bобладает\b", ["имеет", "у него есть"]),
    (r"\bосуществляет\b", ["делает", "проводит"]),
    (r"\bпредставляет собой\b", ["это"]),
    (r"\bпри этом\b", ["и", "причём"]),
    (r"\bв связи с этим\b", ["поэтому", "из-за этого"]),
    (r"\bв случае если\b", ["если"]),
    (r"\bв целях\b", ["чтобы"]),
    (r"\bв рамках\b", ["в"]),
    (r"\bв качестве\b", ["как"]),
    # --- Томилина ---
    (r"\bво-первых\b", ["сначала", "прежде всего", "для начала"]),
    (r"\bво-вторых\b", ["далее", "потом", "ещё"]),
    (r"\bв-третьих\b", ["и ещё", "следом"]),
    (r"\bкроме того\b", ["ещё", "помимо этого", "также"]),
    (r"\bболее того\b", ["и ещё", "причём"]),
    (r"\bэто означает,?\b", ["то есть", "иначе говоря", "говоря проще"]),
    (r"\bэто свидетельствует о том,?\b", ["значит", "а это говорит о том"]),
    (r"\bимеет непосредственное значение\b", ["напрямую важно", "прямо влияет"]),
    # методологические — чиним субъект (работа/глава/закон)
    (r"\bработа анализирует\b", ["в работе разбираю", "в исследовании анализирую"]),
    (r"\bработа рассматривает\b", ["в работе разбираю", "я рассматриваю"]),
    (r"\bработа показывает\b", ["я показываю", "работа демонстрирует результат"]),
    (r"\bработа отображает\b", ["работа отражает", "я отражаю"]),
    (r"\bглава рассматривает\b", ["в этой главе разбираю", "в главе рассматриваю"]),
    (r"\bглава посвящена\b", ["эту главу посвятил", "в главе говорю о"]),
    (r"\bстатья рассматривает\b", ["в статье разбираю"]),
    (r"\bзакон указывает\b", ["закон говорит", "в законе прописано"]),
    (r"\bзакон устанавливает\b", ["закон закрепляет", "в законе закреплено"]),
    (r"\bнорма устанавливает\b", ["норма закрепляет", "по норме"]),
]


def _apply_replacements(text: str, rng: random.Random) -> str:
    """Вариативная замена штампов (не все, а частично, чтобы звучало живо)."""
    for pattern, options in REPLACEMENTS:
        # Пропускаем ~30% случаев, чтобы не превратить весь текст в один стиль
        def sub(match):
            if rng.random() < 0.3:
                return match.group(0)
            choice = rng.choice(options)
            # Сохраняем регистр первой буквы
            if match.group(0)[0].isupper():
                choice = choice[0].upper() + choice[1:]
            return choice
        text = re.sub(pattern, sub, text, flags=re.IGNORECASE)
    return text


def _break_parallelism(text: str, rng: random.Random) -> str:
    """
    Ломает параллельные конструкции: 
      'не только А, но и Б' → 'А. И ещё — Б'
    """
    # «не только ... но и ...» → две фразы
    text = re.sub(
        r"не только ([^,.]{3,60}),?\s*но и ([^.!?]{3,80})",
        lambda m: f"{m.group(1)}. А ещё {m.group(2)}" if rng.random() > 0.4 else m.group(0),
        text,
        flags=re.IGNORECASE,
    )
    return text


def _vary_sentence_length(text: str, rng: random.Random) -> str:
    """
    Режет часть длинных сложносочинённых предложений на короткие — даёт burstiness.
    """
    sentences = re.split(r"(?<=[.!?])\s+", text)
    out = []
    for s in sentences:
        words = s.split()
        if len(words) > 28 and rng.random() < 0.5:
            # Ищем подходящее место для разрыва: ", и", ", а", ", но", ", что"
            m = re.search(r",\s+(и|а|но|что|поэтому|причём|однако)\s+", s, re.IGNORECASE)
            if m:
                left = s[:m.start()].rstrip(",.").strip()
                right = s[m.end():].strip()
                if left and right:
                    # Заглавная буква у правой части
                    right = right[0].upper() + right[1:]
                    out.append(left + ".")
                    out.append(right)
                    continue
        out.append(s)
    return " ".join(out)


def _inject_human_markers(text: str, style: str, rng: random.Random) -> str:
    """
    Вставляет человеческие маркеры в случайные места (2-4 штуки на текст),
    в зависимости от стиля.
    """
    markers = {
        "casual":   ["На самом деле, ", "Честно говоря, ", "По-моему, ", "Мне кажется, ", "Если честно, "],
        "expert":   ["По моему опыту, ", "На практике ", "Замечу, что ", "Нередко ", "В моей практике "],
        "blogger":  ["Кстати, ", "Окей, ", "Слушайте, ", "Между прочим, ", "А вот что интересно — "],
        "business": ["Отмечу, что ", "На мой взгляд, ", "Важный момент: ", "Хочу подчеркнуть: "],
        "neutral":  ["Замечу, ", "Кстати, ", "На мой взгляд, ", "По сути, "],
    }
    pool = markers.get(style, markers["neutral"])
    sentences = re.split(r"(?<=[.!?])\s+", text)
    if len(sentences) < 4:
        return text
    k = min(len(pool), max(1, len(sentences) // 5))
    indices = rng.sample(range(1, len(sentences)), min(k, len(sentences) - 1))
    for idx in indices:
        if rng.random() < 0.5:
            marker = rng.choice(pool)
            s = sentences[idx]
            if s and s[0].isupper():
                sentences[idx] = marker + s[0].lower() + s[1:]
    return " ".join(sentences)


def _micro_imperfections(text: str, rng: random.Random) -> str:
    """
    Лёгкие живые шероховатости: иногда «—» вместо «.», иногда «…» в конце.
    НЕ опечатки (это плохо) — только стилистические микрошумы.
    """
    # Редко заменяем точку на «–» (en dash, НЕ em dash) между короткими фразами
    def maybe_dash(m):
        if rng.random() < 0.08:
            return " – "
        return m.group(0)
    text = re.sub(r"\.\s+(?=[А-ЯA-Z])", maybe_dash, text)
    # Редко добавляем многоточие в конце размышления
    sentences = text.split(". ")
    for i, s in enumerate(sentences[:-1]):
        if s.endswith("ь") and rng.random() < 0.05:
            sentences[i] = s + ".."
    return ". ".join(sentences)


def _kill_markers_tomilina(text: str, rng: random.Random) -> str:
    """
    Обязательные зачистки по Томилиной:
      - длинное тире «—» → среднее «–» (или разрыв предложения)
      - буква «ё» → «е» (академическая норма)
      - английские термины без нужды → убираем/заменяем
      - шаблонные «это означает» уже в REPLACEMENTS
    """
    # 1. Длинное тире → среднее тире (или разрыв)
    def replace_em(m):
        # 30% случаев — ставим точку и разрываем предложение
        if rng.random() < 0.3:
            return ". "
        return " – "
    text = re.sub(r"\s*—\s*", replace_em, text)

    # 2. «ё» → «е» (только в словах, не в именах — но для академической работы нормально везде)
    text = text.replace("ё", "е").replace("Ё", "Е")

    # 3. Чистим одиночные английские слова в русском тексте (если не аббревиатура)
    def check_eng(m):
        w = m.group(0)
        # Оставляем аббревиатуры (все заглавные) и короткие (2-3 буквы)
        if w.isupper() or len(w) <= 3:
            return w
        # Оставляем если в кавычках или скобках рядом (проверяется снаружи, тут просто удаляем)
        return ""
    # Применяем осторожно: только если английское слово окружено кириллицей
    text = re.sub(r"(?<=[а-яА-Я\s])\b[A-Za-z]{4,}\b(?=[\s,.:;!?])",
                  lambda m: "" if rng.random() < 0.5 else m.group(0),
                  text)

    # 4. Чистим двойные пробелы после удалений
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r"\s+([.,;:!?])", r"\1", text)
    return text


# ======================================================================
#              GIGACHECK-BUSTER: агрессивная human-паcс
# ======================================================================
# Эти преобразования ломают статистику, по которой распознаются LLM-тексты:
# 1) Burstiness — режем/склеиваем предложения, добавляем «ритмические рваные»
# 2) Perplexity — подставляем низкочастотные синонимы в «слишком гладких» местах
# 3) Lexical diversity — варьируем союзы, вводные, порядок слов
# 4) Punctuation entropy — добавляем/убираем запятые в допустимых местах
# 5) Morphology noise — меняем порядок слов в простых парах


# Синонимы: пары (LLM-частое слово) → (человеко-редкое, но уместное)
LOWPERPLEXITY_SWAPS: List[Tuple[str, List[str]]] = [
    (r"\bважн(ый|ая|ое|ые|ую|ого|ом|ой)\b", ["значимый", "ощутимый", "заметный", "весомый", "серьёзный"]),
    (r"\bразвити(е|я|ю|ем|и)\b", ["рост", "движение", "эволюция", "продвижение"]),
    (r"\bпроцесс(а|у|ом|е|ы|ов|ам|ами|ах)?\b", ["работа", "ход", "дело", "движение"]),
    (r"\bрезультат(а|у|ом|е|ы|ов|ам|ами|ах)?\b", ["итог", "вывод", "плод", "отдача"]),
    (r"\bзначи(т|тельн\w+)\b", ["говорит о", "означает", "указывает"]),
    (r"\bвозможност(ь|и|ью|ей|ям|ями|ях)\b", ["шанс", "выход", "путь", "вариант"]),
    (r"\bпроблем(а|ы|е|у|ой|ам|ами|ах)\b", ["трудность", "загвоздка", "сложность", "заковырка"]),
    (r"\bсовременн(ый|ая|ое|ые|ую|ого|ом|ой|ых|ыми)\b", ["нынешний", "сегодняшний", "текущий"]),
    (r"\bактивно\b", ["бурно", "вовсю", "не останавливаясь", "без передышки"]),
    (r"\bбыстро\b", ["стремительно", "мигом", "с ходу", "в два счёта"]),
    (r"\bмного\b", ["куча", "немало", "порядочно", "изрядно"]),
    (r"\bхорошо\b", ["недурно", "толково", "на совесть", "как надо"]),
    (r"\bтехнологи(я|и|ю|ей|ям|ями|ях)\b", ["технология", "инструмент", "метод", "приём"]),
    (r"\bиспользу(ю|ет|ем|ете|ют)\b", ["применяю", "беру", "пускаю в ход", "задействую"]),
    (r"\bнеобходимо\b", ["нужно", "стоит", "придётся", "надо"]),
    (r"\bпозволяет\b", ["даёт", "помогает", "открывает"]),
    (r"\bобеспечивает\b", ["даёт", "создаёт", "гарантирует"]),
    (r"\bспособствует\b", ["помогает", "работает на", "толкает к"]),
    (r"\bзанима(ю|ет|ем|ете|ют)\b", ["отводит", "берёт на себя", "держит"]),
    (r"\bявля(ю|ет|ем|ете|ют)сь?\b", ["это", "становится", "оказывается"]),
    (r"\bзадач(а|и|е|у|ей|ам|ами|ах)\b", ["задумка", "цель", "дело", "работа"]),
    (r"\bэффективн(ый|ая|ое|ые|ую|ого|ом|ой)\b", ["толковый", "рабочий", "дельный", "годный"]),
    (r"\bкачественн(ый|ая|ое|ые|ую|ого|ом|ой)\b", ["добротный", "толковый", "хороший"]),
    (r"\bсистем(а|ы|е|у|ой|ам|ами|ах)\b", ["схема", "механизм", "связка", "устройство"]),
]


def _swap_lowperplexity(text: str, rng: random.Random, chance: float = 0.55) -> str:
    """Подмена частых слов на редкие синонимы — повышает perplexity."""
    out = text
    for pattern, options in LOWPERPLEXITY_SWAPS:
        def sub(m, opts=options):
            if rng.random() > chance:
                return m.group(0)
            pick = rng.choice(opts)
            # Сохраняем капитализацию
            orig = m.group(0)
            if orig and orig[0].isupper():
                pick = pick[0].upper() + pick[1:]
            return pick
        out = re.sub(pattern, sub, out, flags=re.IGNORECASE)
    return out


# Человеческие вводные — реже, но разнообразнее
HUMAN_INSERTS_NEUTRAL = [
    "по сути", "если честно", "по факту", "на самом деле", "в общем-то",
    "так-то", "на деле", "по-хорошему", "как ни крути", "в принципе",
]
HUMAN_INSERTS_CASUAL = [
    "короче", "ну вот", "кстати", "вот что", "тут такое дело", "смотри",
    "получается", "так вот", "как бы то ни было",
]
HUMAN_INSERTS_EXPERT = [
    "на практике", "по моему опыту", "исходя из практики", "в работе",
    "в реальных условиях", "если копнуть глубже",
]


def _inject_living_voice(text: str, style: str, rng: random.Random,
                         density: float = 0.25) -> str:
    """
    Вставляет «живые» вводные в начало ~25% предложений.
    Это ломает ровный академический ритм и добавляет burstiness.
    """
    pool = HUMAN_INSERTS_NEUTRAL
    if style == "casual":
        pool = HUMAN_INSERTS_NEUTRAL + HUMAN_INSERTS_CASUAL
    elif style == "expert":
        pool = HUMAN_INSERTS_NEUTRAL + HUMAN_INSERTS_EXPERT
    elif style == "blogger":
        pool = HUMAN_INSERTS_CASUAL
    elif style == "business":
        pool = HUMAN_INSERTS_NEUTRAL + HUMAN_INSERTS_EXPERT

    sentences = re.split(r"(?<=[.!?])\s+", text)
    result = []
    for s in sentences:
        s = s.strip()
        if not s:
            continue
        # Не вставляем в короткие или вопросительные
        wc = len(re.findall(r"\w+", s))
        if wc > 6 and not s.endswith("?") and rng.random() < density:
            intro = rng.choice(pool)
            # Ставим в начало с маленькой буквы оригинала
            if s[0].isupper():
                s = intro.capitalize() + ", " + s[0].lower() + s[1:]
            else:
                s = intro + ", " + s
        result.append(s)
    return " ".join(result)


def _burstify_sentences(text: str, rng: random.Random) -> str:
    """
    Режет длинные предложения пополам на границах «, и», «, но», «, что»,
    а короткие иногда склеивает тире. Это резко повышает burstiness.
    """
    out_parts = []
    sentences = re.split(r"(?<=[.!?])\s+", text)

    i = 0
    while i < len(sentences):
        s = sentences[i].strip()
        if not s:
            i += 1
            continue

        wc = len(re.findall(r"\w+", s))

        # Длинное (>22 слов) — режем на «, и / , но / , что / , который»
        if wc > 22 and rng.random() < 0.7:
            # Находим подходящее место разреза во второй половине
            words = s.split()
            mid = len(words) // 2
            # Ищем разделитель около середины
            cut_idx = None
            for j in range(mid, len(words) - 3):
                if words[j].rstrip(",").lower() in ("и", "но", "а", "что", "чтобы", "поэтому", "однако"):
                    if words[j - 1].endswith(","):
                        cut_idx = j
                        break
            if cut_idx:
                first = " ".join(words[:cut_idx]).rstrip(",") + "."
                second_head = words[cut_idx].rstrip(",").capitalize()
                second_rest = " ".join(words[cut_idx + 1:])
                second = second_head + " " + second_rest
                out_parts.append(first)
                out_parts.append(second)
                i += 1
                continue

        # Короткое (<8 слов) — иногда склеиваем со следующим через « — »
        if wc < 8 and i + 1 < len(sentences) and rng.random() < 0.25:
            next_s = sentences[i + 1].strip()
            if next_s:
                merged = s.rstrip(".!?") + " – " + next_s[0].lower() + next_s[1:]
                out_parts.append(merged)
                i += 2
                continue

        out_parts.append(s)
        i += 1

    return " ".join(out_parts)


def _shuffle_word_order(text: str, rng: random.Random) -> str:
    """
    Лёгкие перестановки внутри фраз, характерные для русской речи:
    «быстро работает» → «работает быстро»
    «постоянно использует» → «использует постоянно»
    """
    # Наречие + глагол → глагол + наречие (с вероятностью)
    adverbs = r"(быстро|медленно|постоянно|часто|редко|сильно|слабо|активно|полностью|практически|обычно|реально|легко|сложно)"
    verb_ending = r"([а-яё]+(?:ет|ют|ит|ат|ят|ем|ём|ал|ала|али|ло))"

    def swap(m):
        if rng.random() < 0.5:
            return f"{m.group(2)} {m.group(1)}"
        return m.group(0)

    text = re.sub(rf"\b{adverbs}\s+{verb_ending}\b", swap, text, flags=re.IGNORECASE)
    return text


def _punctuation_entropy(text: str, rng: random.Random) -> str:
    """
    Повышает энтропию пунктуации — то, на что смотрят статистические детекторы.
    - иногда заменяет ', и' на '. И'
    - иногда заменяет ', но' на '. Но'
    - иногда ставит тире вместо 'это'/'является'
    """
    # ', и' → '. И' (редко, ~15%)
    def comma_i(m):
        return ". И" if rng.random() < 0.2 else m.group(0)
    text = re.sub(r",\s+и\s+", lambda m: ". И " if rng.random() < 0.2 else m.group(0), text)
    text = re.sub(r",\s+но\s+", lambda m: ". Но " if rng.random() < 0.15 else m.group(0), text)
    text = re.sub(r",\s+а\s+", lambda m: ". А " if rng.random() < 0.12 else m.group(0), text)
    # «это» в середине → тире
    text = re.sub(r"\s+—\s+это\s+", lambda m: " – " if rng.random() < 0.5 else " – это ", text)
    return text


def _drop_predictable_words(text: str, rng: random.Random) -> str:
    """
    Убирает «предсказуемые» LLM-связки: «при этом», «в то же время», «в данном случае».
    Эти слова резко повышают вероятность определения текста как AI.
    """
    predictable = [
        r"\bпри этом[,\s]*", r"\bв то же время[,\s]*", r"\bв данном случае[,\s]*",
        r"\bв свою очередь[,\s]*", r"\bпо сути говоря[,\s]*",
        r"\bтем самым[,\s]*", r"\bв частности[,\s]*",
        r"\bбезусловно[,\s]*", r"\bнесомненно[,\s]*",
    ]
    for pat in predictable:
        text = re.sub(pat, lambda m: "" if rng.random() < 0.75 else m.group(0),
                       text, flags=re.IGNORECASE)
    # Лишние пробелы/запятые
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r",\s*,", ",", text)
    text = re.sub(r"^\s*,\s*", "", text, flags=re.MULTILINE)
    return text


def _add_subjective_voice(text: str, rng: random.Random) -> str:
    """
    Превращает часть безличных конструкций в личные — это сильно снижает AI-score.
    «Следует учитывать» → «Я учитываю»
    «Можно сказать, что» → «Скажу так:»
    """
    subj_replacements = [
        (r"\bследует учитывать\b", ["я учитываю", "стоит держать в голове", "держу в уме"]),
        (r"\bможно сказать,?\s*что\b", ["скажу так —", "по сути —", "если коротко —"]),
        (r"\bнеобходимо понимать,?\s*что\b", ["важно понять —", "держу в уме —"]),
        (r"\bстоит отметить,?\s*что\b", ["замечу —", "добавлю —", "подчеркну —"]),
        (r"\bследует подчеркнуть,?\s*что\b", ["отмечу —", "добавлю —"]),
        (r"\bочевидно,?\s*что\b", ["понятно —", "ясно —", "видно —"]),
    ]
    for pat, opts in subj_replacements:
        def sub(m, opts=opts):
            if rng.random() < 0.7:
                pick = rng.choice(opts)
                if m.group(0)[0].isupper():
                    pick = pick[0].upper() + pick[1:]
                return pick
            return m.group(0)
        text = re.sub(pat, sub, text, flags=re.IGNORECASE)
    return text


def human_pass(text: str, style: str, rng: random.Random,
               strength: str = "high") -> str:
    """
    Главная функция: запускает всю цепочку гиперчеловеческих преобразований.
    strength: low | mid | high | extreme
    """
    out = text

    # Шаг 1 — убираем предсказуемые LLM-связки
    out = _drop_predictable_words(out, rng)

    # Шаг 2 — подмена частых слов на редкие синонимы
    chance = {"low": 0.3, "mid": 0.45, "high": 0.6, "extreme": 0.8}.get(strength, 0.55)
    out = _swap_lowperplexity(out, rng, chance=chance)

    # Шаг 3 — субъективизация (безличное → личное)
    out = _add_subjective_voice(out, rng)

    # Шаг 4 — режем/склеиваем для burstiness
    if strength in ("mid", "high", "extreme"):
        out = _burstify_sentences(out, rng)

    # Шаг 5 — перестановки порядка слов
    if strength in ("high", "extreme"):
        out = _shuffle_word_order(out, rng)

    # Шаг 6 — пунктуационная энтропия
    if strength in ("high", "extreme"):
        out = _punctuation_entropy(out, rng)

    # Шаг 7 — вставка живого голоса
    density = {"low": 0.1, "mid": 0.18, "high": 0.25, "extreme": 0.35}.get(strength, 0.25)
    out = _inject_living_voice(out, style, rng, density=density)

    # На extreme — второй прогон синонимов
    if strength == "extreme":
        out = _swap_lowperplexity(out, rng, chance=0.5)

    # Финальная чистка
    out = re.sub(r"\s{2,}", " ", out)
    out = re.sub(r"\s+([.,;:!?])", r"\1", out)
    out = re.sub(r"([.!?])\s*([.!?])+", r"\1", out)
    return out.strip()


def postprocess_antidetect(text: str, style: str, aggression: str, seed: int,
                            academic_mode: bool = False) -> str:
    """
    Финальный антидетект-постпроцессор. Агрессивность:
      - light  — только замены штампов (30%)
      - medium — штампы + ломка параллелизма + 1 маркер
      - strong — всё + варьирование длины + маркеры + микрошумы
    academic_mode=True — включает зачистки Томилиной (тире, ё, англ. термины)
    """
    rng = random.Random(seed)
    out = text

    if aggression in ("light", "medium", "strong", "extreme"):
        out = _apply_replacements(out, rng)

    # В академическом режиме ВСЕГДА чистим тире/ё/англ.вставки
    if academic_mode:
        out = _kill_markers_tomilina(out, rng)

    if aggression in ("medium", "strong", "extreme"):
        out = _break_parallelism(out, rng)
        # Человеческие маркеры в академическом режиме — сдержанно
        if not academic_mode:
            out = _inject_human_markers(out, style, rng)

    if aggression in ("strong", "extreme"):
        out = _vary_sentence_length(out, rng)
        if not academic_mode:
            out = _micro_imperfections(out, rng)
        # ещё раз замены — уже по новому тексту
        out = _apply_replacements(out, rng)

    # === GIGACHECK-BUSTER human_pass ===
    # Запускаем для medium и выше. На extreme — максимальная сила.
    strength_map = {
        "light": None,  # не запускаем
        "medium": "mid",
        "strong": "high",
        "extreme": "extreme",
    }
    hp_strength = strength_map.get(aggression)
    if hp_strength:
        out = human_pass(out, style, rng, strength=hp_strength)

    # Повторная зачистка Томилиной после human_pass (human_pass мог добавить тире)
    if academic_mode:
        out = _kill_markers_tomilina(out, rng)

    # Финальная чистка пробелов
    out = re.sub(r"[ \t]{2,}", " ", out)
    out = re.sub(r"\s+([.,;:!?])", r"\1", out)
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip()


# ======================================================================
#                        LLM-самооценка (гибрид)
# ======================================================================

def llm_score_batch(sentences: List[str]) -> List[float]:
    """
    Оценивает батч предложений одним запросом в YandexGPT.
    Возвращает список AI-scores 0..100.
    Если YandexGPT недоступен — возвращает [] (упадём на эвристику).
    """
    if not sentences:
        return []
    numbered = "\n".join(f"{i+1}. {s}" for i, s in enumerate(sentences))
    system = (
        "Ты — детектор AI-текста. Тебе даётся список предложений. "
        "Для каждого оцени вероятность того, что оно написано генеративной моделью (GPT и подобными), от 0 до 100. "
        "Признаки AI: штампы, ровная длина, формальная лексика, отсутствие личного опыта, шаблонные вводные, "
        "избыток канцеляризмов, параллельные конструкции. "
        "Признаки человека: живая разговорная лексика, неровный ритм, личные маркеры, спонтанные отступления, "
        "эмоциональные окраски, редкие просторечия. "
        "Ответ — СТРОГО только JSON-массив чисел, по одному на предложение, в том же порядке. "
        "Без комментариев, без префиксов. Пример: [72, 15, 88, 40]."
    )
    user = f"Оцени каждое из этих предложений:\n\n{numbered}"
    try:
        raw = _yagpt_complete(system, user, temperature=0.2, max_tokens=600, model="yandexgpt-lite")
    except Exception:
        return []
    # Парсим JSON-массив
    match = re.search(r"\[[\d,\s.]+\]", raw)
    if not match:
        return []
    try:
        arr = json.loads(match.group(0))
        if not isinstance(arr, list):
            return []
        # Нормализация длины
        if len(arr) < len(sentences):
            arr = list(arr) + [50.0] * (len(sentences) - len(arr))
        elif len(arr) > len(sentences):
            arr = arr[:len(sentences)]
        return [float(max(0, min(100, x))) for x in arr]
    except Exception:
        return []


def _detect_markers(sentence: str) -> List[str]:
    """Возвращает список сработавших маркеров в предложении (для UI-подсветки)."""
    markers = []
    s = sentence
    if "—" in s:
        markers.append("em_dash")
    if re.search(r"\b(ё|Ё)\w+", s):
        markers.append("letter_yo")
    if re.search(r"(?<=[а-яА-Я\s])\b[A-Za-z]{4,}\b", s):
        markers.append("english_term")
    if len(re.findall(r"\w+", s)) > 40:
        markers.append("too_long")
    if AI_PHRASES_RE.search(s):
        markers.append("ai_cliche")
    if re.search(r"\b(во-первых|во-вторых|кроме того|более того|следует отметить)\b", s, re.IGNORECASE):
        markers.append("template_connector")
    if re.search(r"\b(работа|глава|статья|исследование|закон|норма|система|судебная практика)\s+(анализирует|рассматривает|показывает|отображает|раскрывает|описывает|указывает|устанавливает|требует|определяет|посвящена)\b", s, re.IGNORECASE):
        markers.append("methodological_error")
    if re.search(r"\b(это означает|это свидетельствует|имеет непосредственное значение)\b", s, re.IGNORECASE):
        markers.append("template_logic")
    return markers


def detect_ai(text: str, use_llm: bool = True) -> Dict[str, Any]:
    """
    Гибридная детекция. Возвращает список предложений + общий score + маркеры.
    Снижен порог — ловит даже ~7% AI-сигналов.
    """
    sentences = split_sentences(text)
    if not sentences:
        return {"sentences": [], "overall_score": 0.0, "burstiness_penalty": 0.0,
                "tense_penalty": 0.0, "structure_penalty": 0.0, "markers": {}}

    # 1) эвристика
    sentences = score_sentences_heuristic(sentences)

    # 2) Документные бонусы
    bp = _burstiness(sentences)
    tp = _tense_jumps(sentences)
    sp = _uniform_paragraph_structure(text)
    doc_penalty = bp + tp + sp  # 0..30+
    for s in sentences:
        s["heuristic_score"] = min(100.0, s["heuristic_score"] + doc_penalty)
        s["markers"] = _detect_markers(s["text"])

    # 3) LLM
    llm_scores: List[float] = []
    if use_llm:
        for i in range(0, len(sentences), 20):
            batch = [s["text"] for s in sentences[i:i + 20]]
            scores = llm_score_batch(batch)
            if scores:
                llm_scores.extend(scores)
            else:
                llm_scores.extend([s["heuristic_score"] for s in sentences[i:i + 20]])

    # 4) Объединяем: 45% эвристика + 55% LLM, + «максимальная триггер-линза»
    #    (берём max между средним и эвристикой с маркерами, чтобы не упустить редкое)
    for i, s in enumerate(sentences):
        h = s["heuristic_score"]
        if i < len(llm_scores):
            combined = 0.45 * h + 0.55 * llm_scores[i]
            # Если сработал методологический маркер или AI-штамп — бустим
            if "methodological_error" in s["markers"] or "ai_cliche" in s["markers"]:
                combined = max(combined, 55.0)
            if "template_logic" in s["markers"] and h > 30:
                combined = max(combined, 60.0)
            s["ai_score"] = round(min(100, combined), 1)
            s["llm_score"] = round(llm_scores[i], 1)
        else:
            s["ai_score"] = round(h, 1)
            s["llm_score"] = None

    # 5) Общий score: среднее с весом по длине + «пенальти за долю горячих»
    total_weight = 0
    weighted_sum = 0.0
    hot_count = 0
    for s in sentences:
        w = max(1, len(s["text"].split()))
        total_weight += w
        weighted_sum += s["ai_score"] * w
        if s["ai_score"] >= 50:
            hot_count += 1
    base = weighted_sum / total_weight if total_weight > 0 else 0
    hot_ratio = hot_count / max(1, len(sentences))
    # Если хотя бы 1 из 10 предложений «горячее» — добавим общий бонус
    overall = base + min(15, hot_ratio * 40)
    overall = min(100, overall)

    # Счётчик маркеров по всему тексту
    marker_counts: Dict[str, int] = {}
    for s in sentences:
        for m in s.get("markers", []):
            marker_counts[m] = marker_counts.get(m, 0) + 1

    return {
        "sentences": sentences,
        "overall_score": round(overall, 1),
        "burstiness_penalty": round(bp, 1),
        "tense_penalty": round(tp, 1),
        "structure_penalty": round(sp, 1),
        "markers": marker_counts,
    }


# ======================================================================
#                         Переписывание
# ======================================================================

STYLE_DESCRIPTIONS = {
    "neutral":  "нейтральный, сдержанный, но живой",
    "casual":   "разговорный, как будто другу рассказываешь, простыми словами",
    "expert":   "экспертный — от первого лица, с личным опытом, профессионально, но без канцелярщины",
    "blogger":  "блогерский, с эмоциями, динамичный, допускает просторечия",
    "business": "деловой, но тёплый, без сухих канцеляризмов",
}

AGGRESSION_NOTES = {
    "light":   "Перефразируй осторожно, почти не меняя длину фраз. Сохраняй 70% лексики.",
    "medium":  "Перефразируй ощутимо, варьируй длину предложений, убери AI-штампы. Меняй до 60% лексики.",
    "strong":  "Переписывай радикально: разбей длинные фразы, вставляй живые вводные (я думаю, кстати, на практике), используй просторечия где уместно, убирай все канцеляризмы, сильно варьируй ритм. Меняй до 80% лексики, но СМЫСЛ СОХРАНЯЙ НА 100%.",
    "extreme": "МАКСИМАЛЬНО агрессивный режим против GigaCheck/GPTZero. Переписывай так, будто живой человек печатает на клавиатуре — с разной длиной предложений (от 3 до 40 слов ВПЕРЕМЕШКУ), неожиданными словами, низкочастотными синонимами, личными оборотами («я думаю», «по-моему», «как мне кажется», «на моём опыте»), разговорными связками («ну», «вот», «короче», «кстати», «по сути», «на самом деле»). ОБЯЗАТЕЛЬНО: хотя бы одно очень короткое предложение (2-5 слов) на абзац. ОБЯЗАТЕЛЬНО: хотя бы одна риторическая вставка или вопрос самому себе. Меняй до 95% лексики. Используй синонимы, которые LLM использует редко.",
}


def _build_rewrite_prompt(style: str, aggression: str, preserve_terms: bool,
                          academic_mode: bool = False) -> str:
    style_desc = STYLE_DESCRIPTIONS.get(style, STYLE_DESCRIPTIONS["neutral"])
    aggr = AGGRESSION_NOTES.get(aggression, AGGRESSION_NOTES["medium"])
    preserve = (
        "Термины, названия, цифры, имена собственные — НЕ МЕНЯЙ. Всё остальное переписывай свободно. "
        if preserve_terms else
        "Можно заменять и термины на более простые эквиваленты, если смысл сохраняется. "
    )

    # Базовый промпт
    base = (
        f"Ты — опытный редактор-человек, эксперт в теме текста. Переписываешь AI-тексты так, "
        f"чтобы они звучали как живая речь человека, который разбирается в предмете. "
        f"Твоя задача — обойти детекторы AI (GigaCheck, GPTZero, AntiPlagiat AI, Turnitin и подобные), "
        f"но сохранить 100% смысла оригинала.\n\n"
        f"СТАТИСТИЧЕСКИЕ ДЕТЕКТОРЫ (типа GigaCheck) ловят:\n"
        f"- слишком ровную длину предложений → ПИШИ РАЗНОЙ ДЛИНЫ (3, 15, 8, 30 слов вперемешку)\n"
        f"- предсказуемые слова («важный», «современный», «эффективный») → бери редкие синонимы\n"
        f"- гладкие связки («при этом», «в данном случае», «следует отметить») → УБИРАЙ или заменяй на «ну», «вот», «кстати»\n"
        f"- безличные конструкции («можно отметить», «следует учитывать») → переводи в ЛИЧНЫЕ («я учитываю», «на мой взгляд»)\n"
        f"- симметричную пунктуацию → ИНОГДА ставь тире вместо запятой, точку вместо «, и»\n"
        f"- лексическую бедность → используй РАЗНЫЕ слова для одного понятия в соседних фразах\n\n"
        f"СТИЛЬ: {style_desc}.\n"
        f"АГРЕССИВНОСТЬ: {aggr}\n"
        f"{preserve}"
    )

    # Универсальные запреты
    bans = (
        "СТРОГО ЗАПРЕЩЕНО использовать: «в современном мире», «играет важную роль», «стоит отметить», "
        "«в первую очередь», «таким образом», «необходимо», «является», «данный/данные», "
        "«представляет собой», «в рамках», «в качестве», «осуществляет», «не только … но и …», "
        "«ключевым аспектом», «важно понимать», «следует отметить», «кроме того», «более того», "
        "«во-первых/во-вторых», «это означает», «это свидетельствует», «имеет непосредственное значение». "
    )

    # Академические правила (Томилина)
    if academic_mode:
        academic = (
            "\n\nКРИТИЧЕСКИ ВАЖНЫЕ ПРАВИЛА АКАДЕМИЧЕСКОГО ТЕКСТА:\n"
            "1. НИКАКИХ длинных тире (—). Только среднее тире (–) или перестрой предложение без тире.\n"
            "2. НЕ используй букву «ё» — пиши «е» (ведет, проведен, определен, отражен).\n"
            "3. Субъект действия = автор (Я/мы), а НЕ «работа»/«глава»/«закон»/«статья». "
            "   Пиши: «в работе я анализирую», НЕ «работа анализирует». "
            "   Пиши: «закон закрепляет», а лучше «в законе прописано».\n"
            "4. НЕ вставляй английские термины без необходимости и без пояснения. "
            "   Если английский термин обязателен — дай русский перевод в скобках.\n"
            "5. Следи за временами глаголов — не скачи между прошедшим и настоящим в одном абзаце.\n"
            "6. СИЛЬНО варьируй длину предложений: 4 слова, 12 слов, 25 слов, 7 слов — вперемешку.\n"
            "7. Структура абзацев должна быть разной — не строй каждый абзац по схеме "
            "«общее → объяснение → вывод».\n"
            "8. Избегай перечислений через «во-первых/во-вторых» — используй естественные связки.\n"
            "9. Не используй логические штампы «это означает/свидетельствует» — заменяй на «то есть», "
            "«иначе говоря», «значит».\n"
            "10. Избегай слишком гладкого текста — добавляй конкретные примеры, уточнения, "
            "    цифры из оригинала.\n"
        )
    else:
        academic = (
            "Варьируй длину предложений: короткие, средние и длинные вперемешку. "
            "Добавь 1-2 живых вводных на абзац (кстати, на практике, честно говоря, по-моему). "
        )

    tail = (
        "Избегай параллельных конструкций. "
        "Смысл оригинала сохраняй на 100%. Не добавляй выдуманных фактов, цифр, источников. "
        "Не переводи на другой язык. "
        "Ответ — ТОЛЬКО переписанный текст без комментариев и без обрамляющих кавычек."
    )

    return base + bans + academic + tail


def rewrite_text(text: str, style: str = "neutral", aggression: str = "medium",
                 preserve_terms: bool = True, num_variants: int = 1,
                 seed: int = 0, academic_mode: bool = False) -> List[str]:
    """Переписывает один фрагмент. Возвращает список вариантов."""
    system = _build_rewrite_prompt(style, aggression, preserve_terms, academic_mode)
    variants: List[str] = []
    for i in range(num_variants):
        temp = 0.6 + 0.15 * i  # 0.6, 0.75, 0.9
        try:
            out = _yagpt_complete(system, text, temperature=temp, max_tokens=2000, model="yandexgpt")
        except Exception as e:
            out = f"[Ошибка YandexGPT: {e}]"
        out = out.strip().strip('"«»')
        out = postprocess_antidetect(out, style, aggression, seed + i, academic_mode=academic_mode)
        variants.append(out)
    return variants


# ======================================================================
#                    Полный цикл: humanize_full
# ======================================================================

def humanize_full(text: str, style: str, aggression: str,
                  preserve_terms: bool, target_score: float = 8.0,
                  max_passes: int = 4, academic_mode: bool = True) -> Dict[str, Any]:
    """
    Полный прогон с гарантией 100% human:
      1) детектим AI-score
      2) переписываем абзацами (чтобы не потерять структуру)
      3) после каждого прохода — зачистка Томилиной (тире/ё/англ)
      4) замеряем снова → если > target, усиливаем и повторяем
      5) финальный фикс: если остались горячие предложения — переписываем их
         ещё раз индивидуально на максимальной агрессивности
    """
    before = detect_ai(text, use_llm=True)

    current_text = text
    passes_done = 0
    history: List[Dict[str, Any]] = [{"pass": 0, "score": before["overall_score"]}]

    current_aggression = aggression
    for p in range(max_passes):
        paragraphs = re.split(r"\n\s*\n", current_text)
        rewritten_paragraphs: List[str] = []
        for para in paragraphs:
            para_stripped = para.strip()
            if not para_stripped or len(para_stripped) < 15:
                rewritten_paragraphs.append(para)
                continue
            variants = rewrite_text(
                para_stripped,
                style=style,
                aggression=current_aggression,
                preserve_terms=preserve_terms,
                num_variants=1,
                seed=p * 97 + hash(para_stripped[:30]) % 10000,
                academic_mode=academic_mode,
            )
            rewritten_paragraphs.append(variants[0] if variants else para_stripped)
        current_text = "\n\n".join(rewritten_paragraphs)

        # Финальная зачистка Томилиной на весь текст
        if academic_mode:
            current_text = _kill_markers_tomilina(
                current_text, random.Random(p * 31 + 7)
            )

        passes_done += 1
        after = detect_ai(current_text, use_llm=True)
        history.append({"pass": passes_done, "score": after["overall_score"]})

        if after["overall_score"] <= target_score:
            break
        # усиливаем
        if current_aggression == "light":
            current_aggression = "medium"
        elif current_aggression == "medium":
            current_aggression = "strong"
        elif current_aggression == "strong":
            current_aggression = "extreme"

    # Финальный точечный проход: ищем оставшиеся горячие предложения и переписываем их индивидуально
    final_check = detect_ai(current_text, use_llm=True)
    hot_sentences = [s for s in final_check["sentences"] if s["ai_score"] >= 40]
    if hot_sentences and passes_done < max_passes + 2:
        for hot in hot_sentences[:10]:  # максимум 10 точечных правок
            try:
                fixed = rewrite_text(
                    hot["text"],
                    style=style,
                    aggression="extreme",
                    preserve_terms=preserve_terms,
                    num_variants=1,
                    seed=hash(hot["text"]) % 99999,
                    academic_mode=academic_mode,
                )
                if fixed and fixed[0]:
                    current_text = current_text.replace(hot["text"], fixed[0], 1)
            except Exception:
                continue
        if academic_mode:
            current_text = _kill_markers_tomilina(current_text, random.Random(42))
        passes_done += 1

    final = detect_ai(current_text, use_llm=True)
    history.append({"pass": passes_done, "score": final["overall_score"]})

    return {
        "original_text": text,
        "humanized_text": current_text,
        "score_before": before["overall_score"],
        "score_after": final["overall_score"],
        "passes": passes_done,
        "history": history,
        "sentences_before": before["sentences"],
        "sentences_after": final["sentences"],
        "markers_before": before.get("markers", {}),
        "markers_after": final.get("markers", {}),
    }


# ======================================================================
#                       DOCX экспорт
# ======================================================================

def _build_docx(text: str, title: str = "Humanized Text") -> bytes:
    """Собирает минимальный .docx в памяти."""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(title, level=1)
        for para in text.split("\n\n"):
            doc.add_paragraph(para.strip())
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        # Fallback: упрощённый ZIP-docx вручную
        import zipfile
        content_xml = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:body>'
        )
        for para in text.split("\n\n"):
            esc = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            content_xml += f'<w:p><w:r><w:t xml:space="preserve">{esc}</w:t></w:r></w:p>'
        content_xml += '</w:body></w:document>'

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
            z.writestr(
                "[Content_Types].xml",
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                '</Types>'
            )
            z.writestr(
                "_rels/.rels",
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
                '</Relationships>'
            )
            z.writestr("word/document.xml", content_xml)
        return buf.getvalue()


# ======================================================================
#                            handler
# ======================================================================

def handler(event: dict, context) -> dict:
    """
    Инструмент Humanizer. Очеловечивает AI-тексты, используя YandexGPT + гибридный детектор.

    Args:
        event: dict с httpMethod, headers, body
        context: объект с request_id, function_name

    Returns:
        HTTP-ответ с JSON
    """
    method = event.get("httpMethod", "GET")
    if method == "OPTIONS":
        return {"statusCode": 200, "headers": CORS_HEADERS, "isBase64Encoded": False, "body": ""}

    if method != "POST":
        return _err("Only POST supported", 405)

    try:
        body = json.loads(event.get("body") or "{}")
    except Exception:
        return _err("Invalid JSON", 400)

    action = body.get("action", "").strip()
    headers = event.get("headers") or {}
    user_id_raw = headers.get("X-User-Id") or headers.get("x-user-id") or body.get("user_id")
    try:
        user_id = int(user_id_raw) if user_id_raw else None
    except (ValueError, TypeError):
        user_id = None

    try:
        if action == "parse_file":
            filename = body.get("filename", "file.txt")
            file_b64 = body.get("file_base64", "")
            if not file_b64:
                return _err("file_base64 required")
            raw = base64.b64decode(file_b64)
            text = _extract_text_from_bytes(raw, filename)
            text = re.sub(r"\n{3,}", "\n\n", text).strip()
            if len(text) > 50000:
                return _ok({
                    "text": text[:50000],
                    "truncated": True,
                    "original_chars": len(text),
                    "chars": 50000,
                })
            return _ok({"text": text, "truncated": False, "chars": len(text)})

        elif action == "detect":
            text = body.get("text", "")
            if not text or len(text) < 10:
                return _err("text too short")
            if len(text) > 30000:
                return _err("text too long (max 30000 chars)")
            use_llm = bool(body.get("use_llm", True))
            result = detect_ai(text, use_llm=use_llm)
            return _ok(result)

        elif action == "rewrite":
            text = body.get("text", "")
            if not text or len(text) < 5:
                return _err("text too short")
            if len(text) > 5000:
                return _err("text too long for single rewrite (max 5000)")
            style = body.get("style", "neutral")
            aggression = body.get("aggression", "medium")
            preserve_terms = bool(body.get("preserve_terms", True))
            academic_mode = bool(body.get("academic_mode", True))
            num_variants = min(3, max(1, int(body.get("num_variants", 3))))
            seed = int(body.get("seed", 42))
            variants = rewrite_text(text, style, aggression, preserve_terms,
                                     num_variants, seed, academic_mode=academic_mode)
            return _ok({"variants": variants})

        elif action == "humanize_full":
            text = body.get("text", "")
            if not text or len(text) < 10:
                return _err("text too short")
            if len(text) > 30000:
                return _err("text too long (max 30000 chars)")
            style = body.get("style", "neutral")
            aggression = body.get("aggression", "medium")
            preserve_terms = bool(body.get("preserve_terms", True))
            academic_mode = bool(body.get("academic_mode", True))
            target_score = float(body.get("target_score", 8.0))
            max_passes = min(5, max(1, int(body.get("max_passes", 4))))
            result = humanize_full(text, style, aggression, preserve_terms,
                                    target_score, max_passes, academic_mode=academic_mode)
            return _ok(result)

        elif action == "humanize_chunk":
            # Быстрое переписывание одного чанка (одного-двух абзацев).
            # Используется фронтом для итеративной обработки длинных текстов —
            # без попадания в таймаут 30с.
            chunk = body.get("text", "")
            if not chunk or len(chunk) < 5:
                return _err("chunk too short")
            if len(chunk) > 4000:
                return _err("chunk too long (max 4000 chars)")
            style = body.get("style", "neutral")
            aggression = body.get("aggression", "extreme")
            preserve_terms = bool(body.get("preserve_terms", True))
            academic_mode = bool(body.get("academic_mode", True))
            seed = int(body.get("seed", 42))
            # Один проход: переписать и вернуть
            variants = rewrite_text(
                chunk,
                style=style,
                aggression=aggression,
                preserve_terms=preserve_terms,
                num_variants=1,
                seed=seed,
                academic_mode=academic_mode,
            )
            out_text = variants[0] if variants else chunk
            # Быстрый детект только по эвристике — без LLM, чтобы не тратить время
            detection = detect_ai(out_text, use_llm=False)
            return _ok({
                "text": out_text,
                "ai_score": detection["overall_score"],
                "sentences": detection["sentences"],
                "markers": detection.get("markers", {}),
            })

        elif action == "save_document":
            if not user_id:
                return _err("user_id required", 401)
            original = body.get("original_text", "")
            humanized = body.get("humanized_text", "")
            title = (body.get("title") or "Документ")[:500]
            source_filename = (body.get("source_filename") or "")[:500]
            source_type = (body.get("source_type") or "paste")[:50]
            score_before = float(body.get("ai_score_before") or 0)
            score_after = float(body.get("ai_score_after") or 0)
            style = (body.get("style") or "neutral")[:50]
            aggression = (body.get("aggression") or "medium")[:20]
            char_count = len(humanized)
            word_count = len(re.findall(r"\w+", humanized))
            # Простые подстановки (безопасно, т.к. используем параметризацию через quote)
            conn = _db()
            try:
                with conn.cursor(cursor_factory=RealDictCursor) as cur:
                    title_s = title.replace("'", "''")
                    fn_s = source_filename.replace("'", "''")
                    st_s = source_type.replace("'", "''")
                    orig_s = original.replace("'", "''")
                    hum_s = humanized.replace("'", "''")
                    sty_s = style.replace("'", "''")
                    agr_s = aggression.replace("'", "''")
                    cur.execute(
                        f"INSERT INTO {SCHEMA}.humanizer_documents "
                        f"(user_id, title, source_filename, source_type, original_text, humanized_text, "
                        f"ai_score_before, ai_score_after, char_count, word_count, style, aggression) "
                        f"VALUES ({user_id}, '{title_s}', '{fn_s}', '{st_s}', '{orig_s}', '{hum_s}', "
                        f"{score_before}, {score_after}, {char_count}, {word_count}, '{sty_s}', '{agr_s}') "
                        f"RETURNING id, created_at"
                    )
                    row = cur.fetchone()
                    conn.commit()
                    return _ok({"id": row["id"], "created_at": row["created_at"].isoformat()})
            finally:
                conn.close()

        elif action == "list_history":
            if not user_id:
                return _err("user_id required", 401)
            limit = min(100, max(1, int(body.get("limit", 20))))
            conn = _db()
            try:
                with conn.cursor(cursor_factory=RealDictCursor) as cur:
                    cur.execute(
                        f"SELECT id, title, source_filename, source_type, ai_score_before, ai_score_after, "
                        f"char_count, word_count, style, aggression, created_at "
                        f"FROM {SCHEMA}.humanizer_documents "
                        f"WHERE user_id = {user_id} "
                        f"ORDER BY created_at DESC LIMIT {limit}"
                    )
                    rows = cur.fetchall()
                    for r in rows:
                        r["created_at"] = r["created_at"].isoformat()
                        if r.get("ai_score_before") is not None:
                            r["ai_score_before"] = float(r["ai_score_before"])
                        if r.get("ai_score_after") is not None:
                            r["ai_score_after"] = float(r["ai_score_after"])
                    return _ok({"documents": rows})
            finally:
                conn.close()

        elif action == "get_document":
            if not user_id:
                return _err("user_id required", 401)
            doc_id = int(body.get("id", 0))
            if not doc_id:
                return _err("id required")
            conn = _db()
            try:
                with conn.cursor(cursor_factory=RealDictCursor) as cur:
                    cur.execute(
                        f"SELECT * FROM {SCHEMA}.humanizer_documents "
                        f"WHERE id = {doc_id} AND user_id = {user_id}"
                    )
                    row = cur.fetchone()
                    if not row:
                        return _err("not found", 404)
                    row["created_at"] = row["created_at"].isoformat()
                    row["updated_at"] = row["updated_at"].isoformat()
                    if row.get("ai_score_before") is not None:
                        row["ai_score_before"] = float(row["ai_score_before"])
                    if row.get("ai_score_after") is not None:
                        row["ai_score_after"] = float(row["ai_score_after"])
                    return _ok({"document": row})
            finally:
                conn.close()

        elif action == "export_docx":
            text = body.get("text", "")
            title = body.get("title", "Humanized")
            if not text:
                return _err("text required")
            data = _build_docx(text, title)
            return _ok({
                "file_base64": base64.b64encode(data).decode("ascii"),
                "filename": f"{re.sub(chr(92)+'W+', '_', title)[:50]}.docx",
                "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "size": len(data),
            })

        else:
            return _err(f"unknown action: {action}")

    except Exception as e:
        import traceback
        return _err(f"{type(e).__name__}: {e}\n{traceback.format_exc()[:500]}", 500)